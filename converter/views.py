from io import BytesIO
from django.shortcuts import render
from .models import ConvertedFile
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
import os
from django.conf import settings


def convert_pdf_to_doc(request):
    """
    Convert a PDF file to DOCX.

    Parameters:
    - request: The HTTP request object.

    Returns:
    - Rendered HTML page.
    """
    if request.method == 'POST' and request.FILES['pdf_file']:
        pdf_file = request.FILES['pdf_file']
        try:
            converted_file_path = handle_pdf_to_doc_conversion(pdf_file)
            converted_file = ConvertedFile.objects.create(original_file=pdf_file, converted_file=converted_file_path)
            success_msg = f"Successfully converted '{pdf_file.name}' to DOCX."
            return render(request, 'result.html', {'converted_file': converted_file, 'success_msg': success_msg})
        except ValueError as e:
            error_msg = str(e)
            return render(request, 'convert.html', {'error_msg': error_msg})
    return render(request, 'convert.html')

def handle_pdf_to_doc_conversion(pdf_file):
    """
    Handle the conversion of a PDF file to DOCX.

    Parameters:
    - pdf_file: The input PDF file.

    Returns:
    - Path to the converted DOCX file.
    """
    # Check file type
    if not pdf_file.name.lower().endswith('.pdf'):
        raise ValueError("Unsupported file type. Please upload a PDF file.")

    pdf_content = pdf_file.read()
    
    # Construct the output directory based on MEDIA_ROOT
    output_directory = os.path.join(settings.MEDIA_ROOT, 'converted', '')
    os.makedirs(output_directory, exist_ok=True)
    
    safe_filename = os.path.splitext(os.path.basename(pdf_file.name))[0].replace(" ", "_")
    converted_file_path = os.path.join(output_directory, f"{safe_filename}.docx")

    with open(converted_file_path, 'wb') as converted_file:
        pdf_reader = PdfReader(BytesIO(pdf_content))
        doc = Document()
        for page_num in range(len(pdf_reader.pages)):
            text = pdf_reader.pages[page_num].extract_text()
            doc.add_paragraph(text)
        doc.save(converted_file)

    return converted_file_path

def convert_doc_to_pdf(request):
    """
    Convert a DOC file to PDF.

    Parameters:
    - request: The HTTP request object.

    Returns:
    - Rendered HTML page.
    """
    if request.method == 'POST' and request.FILES['doc_file']:
        doc_file = request.FILES['doc_file']
        try:
            converted_file_path = handle_doc_to_pdf_conversion(doc_file)
            converted_file = ConvertedFile.objects.create(original_file=doc_file, converted_file=converted_file_path)
            success_msg = f"Successfully converted '{doc_file.name}' to PDF."
            return render(request, 'result.html', {'converted_file': converted_file, 'success_msg': success_msg})
        except NotImplementedError as e:
            error_msg = str(e)
            return render(request, 'convert_doc_to_pdf.html', {'error_msg': error_msg})
        except ValueError as e:
            error_msg = str(e)
            return render(request, 'convert_doc_to_pdf.html', {'error_msg': error_msg})
    return render(request, 'convert_doc_to_pdf.html')

def handle_doc_to_pdf_conversion(doc_file):
    """
    Handle the conversion of a DOC file to PDF.

    Parameters:
    - doc_file: The input DOC file.

    Returns:
    - Path to the converted PDF file.
    """
    # Check file type
    if not doc_file.name.lower().endswith(('.doc', '.docx')):
        raise ValueError("Unsupported file type. Please upload a DOC or DOCX file.")

    if doc_file.multiple_chunks():
        raise NotImplementedError("Handling large files is not implemented.")

    docx_content = doc_file.read()
    
    # Construct the output directory based on MEDIA_ROOT
    output_directory = os.path.join(settings.MEDIA_ROOT, 'converted', '')
    os.makedirs(output_directory, exist_ok=True)
    
    safe_filename = os.path.splitext(os.path.basename(doc_file.name))[0].replace(" ", "_")
    pdf_file_path = os.path.join(output_directory, f"{safe_filename}.pdf")
    
    create_pdf_from_docx(docx_content, pdf_file_path)

    return pdf_file_path

def create_pdf_from_docx(docx_content, pdf_path):
    """
    Create a PDF from DOCX content.

    Parameters:
    - docx_content: Content of the input DOCX file.
    - pdf_path: Path to save the generated PDF.
    """
    # Load DOCX content
    doc = Document(BytesIO(docx_content))

    # Create a PDF document
    pdf = FPDF()
    pdf.add_page()

    # Iterate through paragraphs in the DOCX and add them to the PDF
    for paragraph in doc.paragraphs:
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, paragraph.text)

    # Save the PDF
    pdf.output(pdf_path)
